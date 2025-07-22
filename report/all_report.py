import json
from pathlib import Path
from typing import Union


try:
    from docx.shared import Inches
except ImportError:
    import freecad_funcs
    package = 'python-docx'
    freecad_funcs.install_package(package)

try:
    import math2docx
except ImportError:
    import freecad_funcs
    package = 'math2docx'
    freecad_funcs.install_package(package)

import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


import math2docx

from report import latex_str as latex
from report import strings

etabs_api_path = Path(__file__).absolute().parent.parent


def add_table_figure(
        doc,
        caption: str,
        type_: str='Table ', # 'Figure '
        ):
    paragraph = doc.add_paragraph(type_, style='Caption')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = run = paragraph.add_run()
    r = run._r
    feild_charachter = OxmlElement('w:fldChar')
    feild_charachter.set(qn('w:fldCharType'), 'begin')
    r.append(feild_charachter)
    instr_text = OxmlElement('w:instrText')
    instr_text.text = f' SEQ {type_}\\* ARABIC'
    r.append(instr_text)
    feild_charachter = OxmlElement('w:fldChar')
    feild_charachter.set(qn('w:fldCharType'), 'end')
    r.append(feild_charachter)
    paragraph.add_run(f': {caption} ')

def add_table_of_content(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    fldChar.set(qn('w:dirty'), 'true')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    doc.add_page_break()
    return doc

def add_json_table_to_doc(
                          json_file: str,
                          doc=None,
                          caption: str='',
                       ):
    # create a new document
    if doc is None:
        doc = create_doc()
    try:
        with open(json_file, 'r') as file:
            json_table = json.load(file)
    except json.JSONDecodeError:
        return doc
    table_style = 'List Table 4 Accent 5'
    if isinstance(json_table, list):
        cols = json_table[-1].get('col') + 1
        rows = json_table[-1].get('row') + 1
        table_docx = doc.add_table(rows=rows, cols=cols)
        table_docx.style = table_style
        # write the the table to doc
        for d in json_table:
            row = d.get('row')
            col = d.get('col')
            text = d.get('text')
            color = d.get('color')
            cell = table_docx.cell(row, col)
            cell.text = text
            if row == 0:  # Set header text to bold
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.italic = True
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "#244061"))
                cell._tc.get_or_add_tcPr().append(shading_elm)
            elif color:   # write the data to the remaining rows of the table
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                cell._tc.get_or_add_tcPr().append(shading_elm)
        add_table_figure(doc=doc, caption=caption, type_='Table ')
    if isinstance(json_table, dict) and "settings" in str(json_file):
        pass

    return doc

def create_doc(
        table_of_content: bool=False,
):
    filepath = etabs_api_path / 'report' / 'templates' / 'beam_deflections.docx'
    doc = docx.Document(str(filepath))
    if table_of_content:
        doc = add_table_of_content(doc)
    return doc

def add_bold_paragraph(doc, text, empty_line=True):
    if empty_line:
        doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_simple_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(rows):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value
        cell_label.paragraphs[0].runs[0].bold = True
    return table

def add_table(doc, cols, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(cols) + 1)  # +1 for caption column
    table.style = 'Table Grid'
    # Add caption row
    cols = [''] + list(cols)  # Add empty caption column
    for j, col in enumerate(cols):
        cell = table.cell(0, j)  # +1 to skip caption column
        cell.text = col
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.italic = True
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "#244061"))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    for i, data in enumerate(rows):
        for j, value in enumerate(data):
            cell = table.cell(i + 1, j)
            cell.text = str(value)
    return table

def add_checkbox_table(doc, items):
    table = doc.add_table(rows=len(items), cols=1)
    table.style = 'Table Grid'
    for i, (label, checked) in enumerate(items):
        cell = table.cell(i, 0)
        checkbox = "✓" if checked else "☐"
        cell.text = f"{checkbox} {label}"
    return table

def add_model_settings_report(doc, json_file):
    with open(json_file, encoding="utf-8") as f:
        data = json.load(f)

    # 1. Project Info
    doc.add_page_break()
    add_bold_paragraph(doc, "مشخصات پروژه", empty_line=False)
    info_rows = [
        ("استان", data.get("ostan", "")),
        ("شهر", data.get("city", "")),
        ("سطح خطر نسبی", data.get("risk_level", "")),
        ("نوع خاک", data.get("soil_type", "")),
        ("ضریب اهمیت", data.get("importance_factor", "")),
    ]
    add_simple_table(doc, [row for row in info_rows if row[1]])


    # 2.1 Loads
    doc.add_page_break()
    add_bold_paragraph(doc, "بارهای وارد بر سازه", empty_line=False)
    loads_rows = [
        ("مرده", data.get("dead_combobox", "")),
        ("مرده کفسازی و تیغه‌بندی", data.get("sdead_combobox", "")),
        ("پارتیشن مرده", data.get("partition_dead_combobox", "")),
        ("زنده غیرقابل کاهش", data.get("live_combobox", "")),
        ("زنده قابل کاهش", data.get("lred_combobox", "")),
        ("زنده پارکینگ", data.get("live_parking_combobox", "")),
        ("پارتیشن زنده", data.get("partition_live_combobox", "")),
        ("زنده با ضریب نیم", data.get("live5_combobox", "")),
        ("زنده بام", data.get("lroof_combobox", "")),
        ("برف", data.get("snow_combobox", "")),
        ("تنظیم جرم", data.get("mass_combobox", "")),
        ("زلزله قائم", data.get("ev_combobox", "")),
        ("مودال", data.get("modal_combobox", "")),
    ]
    add_simple_table(doc, [row for row in loads_rows if row[1]])

    # 2.2 Retainin walls
    if data.get("retaining_wall_groupbox"):
        add_bold_paragraph(doc, "بارهای دیوار حائل")
        loads_rows = [
            ("دیوار حائل در جهت مثبت ایکس", data.get("hxp_combobox", "")),
            ("دیوار حائل در جهت منفی ایکس", data.get("hxn_combobox", "")),
            ("دیوار حائل در جهت مثبت وای", data.get("hyp_combobox", "")),
            ("دیوار حائل در جهت منفی وای", data.get("hyn_combobox", "")),
        ]
        add_simple_table(doc, [row for row in loads_rows if row[1]])

    # 3. Dynamic Loads
    if data.get("dynamic_analysis_groupbox"):
        doc.add_page_break()
        add_bold_paragraph(doc, "لودکیس های دینامیکی", empty_line=False)
        scale_factor_text = f"ضریب همپایه سازی جهت ایکس: {data.get('x_scalefactor_combobox', '')}, ضریب همپایه سازی جهت وای: {data.get('y_scalefactor_combobox', '')}"
        doc.add_paragraph(scale_factor_text)
        # Only one table: 100-30 or angular
        if data.get("combination_response_spectrum_checkbox"):
            dynamic_rows = [
                ("لودکیس طیفی بدون خروج از مرکزیت", data.get("sx_combobox", ""), data.get("sy_combobox", "")),
                ("لودکیس طیفی با خروج از مرکزیت ", data.get("sxe_combobox", ""), data.get("sye_combobox", "")),
                ("لودکیس طیفی دریفت بدون خروج از مرکزیت", data.get("sx_drift_combobox", ""), data.get("sy_drift_combobox", "")),
                ("لودکیس طیفی دریفت با خروج از مرکزیت ", data.get("sxe_drift_combobox", ""), data.get("sye_drift_combobox", "")),
            ]
            cols = ("X", "Y")
            add_table(doc, cols, dynamic_rows)
        elif data.get("angular_response_spectrum_checkbox"):
            doc.add_paragraph("Angular Dynamic Loads Table")
            # You can add angular table here if needed

    # 4. Irregularity
    doc.add_page_break()
    add_bold_paragraph(doc, "نامنظمی‌های سازه", empty_line=False)
    # 4.1 Horizontal Irregularity
    add_bold_paragraph(doc, "نامنظمی‌ در پلان", empty_line=False)
    horizontal_irregularity_items = [
        ("نامنظمی هندسی", data.get("reentrance_corner_checkbox", False)),
        ("نامنظمی در دیافراگم", data.get("diaphragm_discontinuity_checkbox", False)),
        ("نامنظمی خارج از صفحه", data.get("out_of_plane_offset_checkbox", False)),
        ("نامنظمی سیستم های غیرموازی", data.get("nonparallel_system_checkbox", False)),
    ]
    if data.get("torsional_irregularity_groupbox", False):
        if data.get("torsion_irregular_checkbox", False):
            horizontal_irregularity_items.insert(1, ("نامنظمی پیچشی زیاد", True))
        else:
            horizontal_irregularity_items.insert(1, ("نامنظمی پیچشی شدید", True))
    else:
        # If torsional irregularity groupbox is not checked, we assume no torsional irregularity
        horizontal_irregularity_items.insert(0, ("نامنظمی پیچشی", False))
    add_checkbox_table(doc, horizontal_irregularity_items)
    # 4.2 Vertical Irregularity
    add_bold_paragraph(doc, "نامنظمی‌ در ارتفاع")
    vertical_irregularity_items = [
        ("نامنظمی هندسی", data.get("geometric_checkbox", False)),
        ("نامنظمی جرمی", data.get("weight_mass_checkbox", False)),
        ("نامنظمی قطع سیستم باربر جانبی", data.get("in_plane_discontinuity_checkbox", False)),
    ]
    if data.get("stiffness_soft_story_groupbox", False):
        if data.get("stiffness_irregular_checkbox", False):
            vertical_irregularity_items.append(("نامنظمی طبقه نرم", True))
        else:
            vertical_irregularity_items.append(("نامنظمی طبقه خیلی نرم", True))
    else:
        # If torsional irregularity groupbox is not checked, we assume no torsional irregularity
        vertical_irregularity_items.append(("نامنظمی سختی جانبی", False))
    if data.get("lateral_strength_weak_story_groupbox", False):
        if data.get("strength_irregular_checkbox", False):
            vertical_irregularity_items.append(("نامنظمی طبقه ضعیف", True))
        else:
            vertical_irregularity_items.append(("نامنظمی طبقه خیلی ضعیف", True))
    else:
        # If torsional irregularity groupbox is not checked, we assume no torsional irregularity
        vertical_irregularity_items.append(("نامنظمی مقاومت جانبی", False))
    add_checkbox_table(doc, vertical_irregularity_items)

    # 5. First System
    doc.add_page_break()
    if data.get("activate_second_system"):
        add_bold_paragraph(doc, "پارامترهای لرزه‌ای سیستم پایین", empty_line=False)
    else:
        add_bold_paragraph(doc, "پارامترهای لرزه‌ای", empty_line=False)
    first_system_rows = [
        (strings.BOT_STORY_APPLY_EARTHQUAKE, data.get("bot_x_combo", "")),
        (strings.TOP_STORY_APPLY_EARTHQUAKE, data.get("top_x_combo", "")),
        (strings.TOP_STORY_FOR_T, data.get("top_story_for_height", "")),
        (strings.NO_STORIES, str(data.get("no_of_story_x", ""))),
        (strings.HEIGHT_METER, str(data.get("height_x", ""))),
    ]
    if data.get("infill", False):
        first_system_rows.append((strings.INFILL_PANNEL, "بله"))
    else:
        first_system_rows.append((strings.INFILL_PANNEL, "خیر"))
    add_simple_table(doc, [row for row in first_system_rows if row[1]])

    # --- Add EX, EXP, EXN, ... table for first system ---
    ex_table_rows = [
        (strings.EX, data.get("ex_combobox", ""), data.get("ey_combobox", "")),
        (strings.EXP, data.get("exp_combobox", ""), data.get("eyp_combobox", "")),
        (strings.EXN, data.get("exn_combobox", ""), data.get("eyn_combobox", "")),
        (strings.EY, data.get("ex_drift_combobox", ""), data.get("ey_drift_combobox", "")),
        (strings.EYP, data.get("exp_drift_combobox", ""), data.get("eyp_drift_combobox", "")),
        (strings.EYN, data.get("exn_drift_combobox", ""), data.get("eyn_drift_combobox", "")),
    ]
    # Only add rows with at least one value
    ex_table_rows = [row for row in ex_table_rows if row[1] or row[2]]
    if ex_table_rows:
        add_bold_paragraph(doc, "بارهای زلزله استاتیکی")
        add_table(doc, ["X", "Y"], ex_table_rows)

    # 6. Second System (if exists)
    if data.get("activate_second_system"):
        add_bold_paragraph(doc, "پارامترهای لرزه‌ای سیستم بالا")
        second_system_rows = [
            (strings.BOT_STORY_APPLY_EARTHQUAKE, data.get("bot_x1_combo", "")),
            (strings.TOP_STORY_APPLY_EARTHQUAKE, data.get("top_x1_combo", "")),
            (strings.TOP_STORY_FOR_T, data.get("top_story_for_height1", "")),
            (strings.NO_STORIES, str(data.get("no_of_story_x1", ""))),
            (strings.HEIGHT_METER, str(data.get("height_x1", ""))),
        ]
        if data.get("infill_1", False):
            first_system_rows.append((strings.INFILL_PANNEL, "بله"))
        else:
            first_system_rows.append((strings.INFILL_PANNEL, "خیر"))
        if data.get("special_case", False):
            second_system_rows.append((strings.STIFFNESS_10_TIMES, "بله"))
        else:
            second_system_rows.append((strings.STIFFNESS_10_TIMES, "خیر"))
        add_simple_table(doc, [row for row in second_system_rows if row[1]])

        # --- Add EX, EXP, EXN, ... table for second system ---
        ex2_table_rows = [
            (strings.EX, data.get("ex1_combobox", ""), data.get("ey1_combobox", "")),
            (strings.EXP, data.get("exp1_combobox", ""), data.get("eyp1_combobox", "")),
            (strings.EXN, data.get("exn1_combobox", ""), data.get("eyn1_combobox", "")),
            (strings.EY, data.get("ex1_drift_combobox", ""), data.get("ey1_drift_combobox", "")),
            (strings.EYP, data.get("exp1_drift_combobox", ""), data.get("eyp1_drift_combobox", "")),
            (strings.EYN, data.get("exn1_drift_combobox", ""), data.get("eyn1_drift_combobox", "")),
            (strings.EX, data.get("ex1_combobox", ""), data.get("ey1_combobox", "")),
        ]
        ex2_table_rows = [row for row in ex2_table_rows if row[1] or row[2]]
        if ex2_table_rows:
            add_bold_paragraph(doc, "بارهای زلزله استاتیکی سیستم بالا")
            add_table(doc, ["X", "Y"], ex2_table_rows)

    add_earthquake_factor_formulation_section(doc)

def add_earthquake_factor_explanations_section(doc, building):
    """
    Adds a section with all HTML calculation explanations from the building object.
    Each explanation is rendered as formatted text in Persian.
    """
    # List all HTML explanation attributes you want to include
    latexes = building.get_latex()
    doc.add_heading('محاسبه ضریب زلزله', level=1)
    math2docx.add_math(doc.add_paragraph(), r'C = \frac{ABI}{R}')
    math2docx.add_math(doc.add_paragraph(), fr'A = {building.acceleration}, I = {building.importance_factor}')
    # if building
    # for par, latex_text in latexes:
    #     if latex_text:
    #         p = doc.add_paragraph()
    #         math2docx.add_math(p, latex_text)
    return doc

def add_earthquake_factor_formulation_section(doc=None):
    """
    Adds a section with all formula calculation explanations from the building object.
    """
    if doc is None:
        doc = create_doc()
    doc.add_page_break()
    doc.add_heading('محاسبه ضریب زلزله', level=1)
    add_formulation_section(doc, latex.earthquake_formula)
    add_formulation_section(doc, latex.earthquake_b_formula)
    add_formulation_section(doc, latex.earthquake_b1)
    doc.add_paragraph("ضریب اصلاح طیف به شرح زیر تعیین میشود: ")
    doc.add_paragraph("الف- برای پهنه های با خطر نسبی خیلی زیاد و زیاد")
    add_formulation_section(doc, latex.earthquake_n1)
    doc.add_paragraph("ب- برای پهنه های با خطر نسبی متوسط و کم")
    add_formulation_section(doc, latex.earthquake_n2)
    return doc

def add_formulation_section(doc, formulation: str = None):
    """
    Adds a section with the given formulation.
    """
    if formulation:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        math2docx.add_math(p, formulation)
    return doc


# --- integrate into your create_report function ---
def create_report(
    etabs=None,
    filename: str = None,
    doc: 'docx.Document' = None,
    results_path: Union[Path, None]=None,
):
    if results_path is None:
        if etabs is None:
            return
        else:
            name = etabs.get_file_name_without_suffix()
            results_path = etabs.get_filepath() / f"{name}_table_results"
            if filename is None:
                filename = results_path / 'all_reports.docx'
            if not results_path.exists():
                return doc, filename
    if doc is None:
        doc = create_doc()
    for file in results_path.glob('*.json'):
        doc.add_paragraph()
        if "model_settings" in file.name:
            add_model_settings_report(doc, file)
        else:
            caption = ''.join(file.name.split('_'))
            caption = caption.replace('.json', "")
            caption = caption.replace('Model', "")
            caption = caption.replace("Table", "")
            result = ""
            for i, char in enumerate(caption):
                if char.isupper() and i > 0:
                    result += " " + char
                else:
                    result += char
            doc = add_json_table_to_doc(json_file=file, doc=doc, caption=result)
        doc.add_page_break()
    if filename is None:
        filename = results_path / 'all_reports.docx'
    doc.save(filename)
    return doc, filename




