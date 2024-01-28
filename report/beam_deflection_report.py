import math
from pathlib import Path
from decimal import Decimal
import docx
from typing import List
import tempfile

import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.lines import Line2D
try:
    from docx.shared import Inches
except ImportError:
    import freecad_funcs
    package = 'python-docx'
    freecad_funcs.install_package(package)


etabs_api_path = Path(__file__).absolute().parent.parent


def add_lines_to_ax(
        line_coords: List,
        ax,
        edge_color: str='black',
        line_width: float=0.5,
        ):
    '''
    line_coords format : [([x1, x2], [y1, y2]), ...]
    '''
    for coords in line_coords:
        line = Line2D(coords[0], coords[1], linewidth = line_width, color = edge_color)
        ax.add_line(line)
        # p = ax.add_line(coords, edgecolor=edge_color, linewidth=line_width, closed=False)
        # ax.add_patch(p)

def add_rectangles_to_ax(
        rectangle_coords: List,
        ax,
        color: str='black',
        ):
    '''
    rectangle_coords format : [([x1, y1], [x2, y2], [x3, y3], [x4, y4]), ...]
    '''
    for coords in rectangle_coords:
        p = patches.Polygon(coords, edgecolor=color, linewidth=0.5, facecolor=color, closed=True)
        ax.add_patch(p)

def convert_5point_to_8point(cx_, cy_, w_, h_, a_):
    theta = math.radians(a_)
    bbox = np.array([[cx_], [cy_]]) + \
        np.matmul([[math.cos(theta), math.sin(theta)],
                   [-math.sin(theta), math.cos(theta)]],
                  [[-w_ / 2, w_/ 2, w_ / 2, -w_ / 2],
                   [-h_ / 2, -h_ / 2, h_ / 2, h_ / 2]])
    polygon = [(bbox[0][i], bbox[1][i]) for i in range(4)]
    return polygon

def get_beam_columns_coords(etabs, beam_name: str):
    story = etabs.SapModel.FrameObj.GetLabelFromName(beam_name)[1]
    beam_columns = etabs.frame_obj.get_beams_columns_on_stories()
    beams, columns = beam_columns[story]
    beam_coords = []
    for name in beams:
        x1, y1, x2, y2 = etabs.frame_obj.get_xy_of_frame_points(name)
        beam_coords.append(([x1, x2], [y1, y2]))
    polygons = []
    for name in columns:
        x1, y1, x2, y2 = etabs.frame_obj.get_xy_of_frame_points(name)
        rotation = etabs.SapModel.FrameObj.GetLocalAxes(name)[0]
        polygon = convert_5point_to_8point(x1, y1, 50, 50, a_=rotation)
        polygons.append(polygon)
    return beam_coords, polygons

def get_beam_column_ax(
        etabs,
        beam_name: str,
        ) -> Path:
    '''
    Get the axes of beams and columns of story that beam_name belogs to
    '''
    # Get beam, columns coordinates
    beam_coords, column_polygons = get_beam_columns_coords(etabs, beam_name)
    fig, ax = plt.subplots()
    plt.axis('off')
    add_lines_to_ax(beam_coords, ax)
    x1, y1, x2, y2 = etabs.frame_obj.get_xy_of_frame_points(beam_name)
    add_rectangles_to_ax(column_polygons, ax)
    add_lines_to_ax([([x1, x2], [y1, y2])], ax, edge_color='red', line_width=2.0)
    set_ax_boundbox(ax, etabs, beam_name)
    return fig, ax


def set_ax_boundbox(ax, etabs, beam_name, scale=1, bb_add=120):
    units = etabs.get_current_unit()
    story = etabs.SapModel.FrameObj.GetLabelFromName(beam_name)[1]
    x_min, y_min, x_max, y_max = etabs.story.get_story_boundbox(story, len_unit=units[1])
    ax.set_xlim((x_min - bb_add) * scale, (x_max + bb_add) * scale)
    ax.set_ylim((y_min - bb_add) * scale, (y_max + bb_add) * scale)

def get_picture(
        etabs,
        beam_name: str,
        filename='',
        ) -> Path:
    units = etabs.get_current_unit()
    etabs.set_current_unit(units[0], 'cm')
    fig, _ = get_beam_column_ax(etabs, beam_name)
    default_tmp_dir = tempfile._get_default_tempdir()
    if not filename:
        filename = beam_name
    filename_path = Path(default_tmp_dir) / f'{filename}.jpg'
    fig.savefig(str(filename_path), orientation='portrait',
        # papertype='a4',
        bbox_inches='tight', dpi=600)
    plt.clf()
    plt.close(fig)
    etabs.set_current_unit(*units)
    return filename_path

def create_doc():
    filepath = etabs_api_path / 'report' / 'templates' / 'beam_deflections.docx'
    doc = docx.Document(str(filepath))
    return doc

def create_report(
                etabs,
                text_1: str,
                text_2: str,
                beam_name: str,
                filename: str = None,
                doc: 'docx.Document' = None,
                ):
    if doc is None:
        doc = create_doc()
    label, story, _  = etabs.SapModel.FrameObj.GetLabelFromName(beam_name)
    doc.add_heading(f'Beam {beam_name} = {label} on {story}',1)
    image_file_path = get_picture(etabs, beam_name, filename)
    doc.add_picture(str(image_file_path), width=Inches(3.5))
    doc.add_paragraph()
    doc.add_heading(text_1, 4)
    doc.add_paragraph()
    doc.add_heading(text_2, 4)
    if filename:
        doc.save(filename)
    return doc

def create_beam_deflection_report(
        etabs,
        results,
        filename: str,
        doc: 'docx.Document' = None,
    ):
    if not doc:
        doc = create_doc()
    for i in range(len(results[0])):
        def1 = results[0][i]
        def2 = results[1][i]
        text1 = results[2][i]
        

        doc = create_report(etabs, )
        doc.add_page_break()
        progressbar.next(True)
    if filename:
        doc.save(filename)
    progressbar.stop()
    return doc



